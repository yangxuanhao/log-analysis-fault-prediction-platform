# -*- coding: utf-8 -*-
"""
将截图插入到操作手册docx中
"""
import os, glob

from docx import Document
from docx.shared import Inches

MANUAL_DIR = 'manual_screenshots'
DOCX_PATH = '操作手册.docx'
OUTPUT_PATH = '操作手册_补充版.docx'

# 模块配置：(模块名称, 截图序号)
MODULE_MAP = [
    ('日志采集汇聚', '01'),
    ('日志解析引擎', '02'),
    ('异常检测识别', '03'),
    ('日志联机搜索', '04'),
    ('故障预测模型', '05'),
    ('根因定位分析', '06'),
    ('告警管理中心', '07'),
    ('服务拓扑发现', '08'),
    ('智能诊断修复', '09'),
    ('态势感知总览', '10'),
    ('智能报告生成', '11'),
    ('知识图谱构建', '12'),
    ('SLA合规监控', '13'),
    ('模型训练管理', '14'),
    ('系统监控诊断', '15'),
    ('多维可视化', '16'),
    ('告警通知配置', '17'),
    ('自定义仪表盘', '18'),
    ('配置管理中心', '19'),
    ('日志审计归档', '20'),
]

def get_screenshots(module_seq, module_name):
    """获取某个模块的操作截图"""
    pattern = f'*_{module_seq}_{module_name}_*.png'
    files = sorted(glob.glob(os.path.join(MANUAL_DIR, pattern)))
    return files

def insert_image_before(para, image_path, doc, width_inches=6.0):
    """在段落前插入图片"""
    temp_p = doc.add_paragraph()
    temp_p.paragraph_format.alignment = 1  # CENTER
    run = temp_p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    para._element.addprevious(temp_p._element)

def find_caption(doc, module_name):
    """找模块标题段落"""
    for para in doc.paragraphs:
        if module_name in para.text and '模块界面' in para.text:
            return para
    return None

def find_heading(doc, heading_text):
    """找特定标题的段落"""
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading') and heading_text in para.text:
            return para
    return None

def main():
    doc = Document(DOCX_PATH)

    # 1. 登录页面截图 - 放在登录步骤段落之后
    print('=== 登录页面截图 ===')
    login_shot = os.path.join(MANUAL_DIR, '*_01_登录页面.png')
    login_files = sorted(glob.glob(login_shot))
    reg_shot = os.path.join(MANUAL_DIR, '*_02_注册页面.png')
    reg_files = sorted(glob.glob(reg_shot))

    # 找到登录步骤的段落（"步骤4"之后）
    login_para = None
    reg_para = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text.startswith('步骤4') and '验证通过' in text:
            login_para = para
        if text.startswith('步骤4') and '注册' in text:
            reg_para = para

    if login_para and login_files:
        insert_image_before(login_para, login_files[0], doc, 5.5)
        print(f'  登录页面截图已插入')
    else:
        print(f'  ⚠ 登录页面段落未找到: login_para={login_para is not None}, login_files={bool(login_files)}')

    if reg_para and reg_files:
        insert_image_before(reg_para, reg_files[0], doc, 5.5)
        print(f'  注册页面截图已插入')
    else:
        print(f'  ⚠ 注册页面段落未找到')

    # 2. 每个模块截图
    inserted = 0
    for module_name, module_seq in MODULE_MAP:
        caption_para = find_caption(doc, module_name)
        if caption_para is None:
            print(f'⚠ 未找到 {module_name} 的标题段落')
            continue

        screenshots = get_screenshots(module_seq, module_name)
        if not screenshots:
            print(f'⚠ 未找到 {module_name} 的截图')
            continue

        print(f'\n{module_name}: {len(screenshots)} 张截图')
        for screenshot in reversed(screenshots):
            insert_image_before(caption_para, screenshot, doc)
            inserted += 1
            print(f'  + {os.path.basename(screenshot)}')

    # 保存
    doc.save(OUTPUT_PATH)
    print(f'\n✅ 完成！共插入 {inserted} 张操作截图 + 2 张登录截图')
    print(f'保存至: {os.path.abspath(OUTPUT_PATH)}')

if __name__ == '__main__':
    main()
