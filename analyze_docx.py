# -*- coding: utf-8 -*-
"""分析操作手册docx结构的脚本"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

doc = Document('操作手册.docx')

# Analyze document structure
print('=== PARAGRAPHS ===')
for i, para in enumerate(doc.paragraphs):
    style = para.style.name if para.style else 'None'
    text = para.text[:120] if para.text.strip() else '(empty)'
    print(f'[{i:3d}] Style: {style:30s} | {text}')

print()
print('=== IMAGE RELATIONSHIPS ===')
for rel in doc.part.rels.values():
    if 'image' in rel.reltype:
        print(f'  {rel.rId:10s} -> {rel.target_ref}')

# Find which paragraphs contain images
from lxml import etree
nsmap = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
}

print()
print('=== PARAGRAPHS WITH IMAGES ===')
for i, para in enumerate(doc.paragraphs):
    # Check for any drawing elements
    drawings = para._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
    if not drawings:
        drawings = para._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
    if drawings:
        for draw in drawings:
            blip = draw.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
            if blip is not None:
                embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                # Get image size info
                ext = draw.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}ext')
                size_info = ''
                if ext is not None:
                    cx = ext.get('cx')
                    cy = ext.get('cy')
                    size_info = f' size={cx}x{cy}'
                text_before = para.text[:60] if para.text.strip() else '(empty)'
                print(f'  Para [{i:3d}] Style={para.style.name:20s} rel={embed:10s}{size_info} | {text_before}')
