# -*- coding: utf-8 -*-
"""重新分析操作手册docx的完整段落结构"""
import sys; sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r'D:\gzwj\信息系统运行日志智能分析与故障预测及根因定位平台\操作手册.docx')
print(f'总段落数: {len(doc.paragraphs)}')
print()

for i, para in enumerate(doc.paragraphs):
    style = para.style.name
    text = para.text.strip()
    if text:
        print(f'[{i:4d}] {style:25s} | {text[:150]}')
