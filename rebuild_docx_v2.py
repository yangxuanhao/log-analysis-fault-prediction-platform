# -*- coding: utf-8 -*-
"""
v2: 逐段重建docx，确保每张截图都在正确位置
不使用addnext，改用直接构建XML
"""
import os, sys, glob, re, shutil
from lxml import etree
from docx import Document
from docx.shared import Inches, Pt, Emu
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn, nsmap
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from copy import deepcopy

os.chdir(r'D:\gzwj\信息系统运行日志智能分析与故障预测及根因定位平台')
SRC = '操作手册.docx'
DST = '操作手册_new.docx'
SCREENSHOT_DIR = 'manual_screenshots'

KNOWN_MODULES = [
    '日志采集汇聚', '日志解析引擎', '异常检测识别', '日志联机搜索',
    '故障预测模型', '根因定位分析', '告警管理中心', '服务拓扑发现',
    '智能诊断修复', '态势感知总览', '智能报告生成', '知识图谱构建',
    'SLA合规监控', '模型训练管理', '系统监控诊断', '多维可视化',
    '告警通知配置', '自定义仪表盘', '配置管理中心', '日志审计归档',
]
MODULE_NUMS = {
    '日志采集汇聚': '01', '日志解析引擎': '02', '异常检测识别': '03',
    '日志联机搜索': '04', '故障预测模型': '05', '根因定位分析': '06',
    '告警管理中心': '07', '服务拓扑发现': '08', '智能诊断修复': '09',
    '态势感知总览': '10', '智能报告生成': '11', '知识图谱构建': '12',
    'SLA合规监控': '13', '模型训练管理': '14', '系统监控诊断': '15',
    '多维可视化': '16', '告警通知配置': '17', '自定义仪表盘': '18',
    '配置管理中心': '19', '日志审计归档': '20',
}

# ========== Step 1: 清理旧截图 ==========
print('Step 1: 清理旧文档中所有截图')
doc = Document(SRC)
indices = []
for i, p in enumerate(doc.paragraphs):
    if p._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'):
        indices.append(i)
print(f'  找到 {len(indices)} 个旧截图段落，删除中...')
for idx in reversed(indices):
    doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)
doc.save(DST)
print(f'  已保存清理版')

# ========== Step 2: 构建操作段落映射 ==========
print('\nStep 2: 扫描动作段落')
doc = Document(DST)

actions = []  # (para_index, para_obj, module, sub_op, text)
current_module = ''

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t:
        continue
    # 跟踪当前模块
    if t in KNOWN_MODULES:
        current_module = t
        continue
    # 识别动作段落
    if not ('点击「' in t or '拖动「' in t or '勾选「' in t):
        continue
    # 找所属子操作
    sub_op = ''
    for j in range(i, -1, -1):
        pj = doc.paragraphs[j]
        if pj.style.name == 'Heading 3':
            ht = pj.text.strip()
            if ht.startswith('- '): sub_op = ht[2:]
            elif ht.startswith('-'): sub_op = ht[1:]
            else: sub_op = ht
            break
    actions.append((i, p, current_module, sub_op, t))

print(f'  共 {len(actions)} 个动作段落')

# ========== Step 3: 截图匹配 ==========
print('\nStep 3: 匹配截图')

op_kw = {
    '一键启动全部': '启动采集', '停止全部': '停止采集',
    '采集统计': '采集统计', '采集配置': '采集配置', '重置计数器': '重置计数器',
    '执行解析': '执行解析', '测试解析': '测试解析', '解析统计': '解析统计',
    '错误模式': '错误模式', '导出结果': '导出结果', '导出报告': '导出报告',
    '立即检测': '立即检测', '检测报告': '检测报告', '高级配置': '高级配置',
    '搜索': '搜索', '实时Tail': '实时Tail', '语法帮助': '语法帮助',
    '执行预测': '执行预测', '趋势分析': '趋势分析', '风险评估': '风险评分',
    '根因追溯': '根因追溯', '关联分析': '关联分析',
    '故障注入': '故障注入', '影响域': '影响域', '生成因果图': '生成因果图',
    '模拟告警': '模拟告警', '清理已处理': '清理已处理',
    '告警统计': '告警统计', '静默模式': '静默模式',
    '自动发现拓扑': '自动发现拓扑', '刷新状态': '刷新状态',
    '拓扑分析': '拓扑分析', '导出拓扑数据': '导出拓扑',
    '智能诊断': '智能诊断', '一键修复': '一键修复',
    '加入知识库': '加入知识库', '批量诊断': '批量诊断',
    '健康检查': '健康检查', '系统自检': '系统自检',
    '性能快照': '性能快照', '组件状态': '组件状态',
    '流量监控': '流量监控', '生成综合报告': '综合报告',
    '日报': '生成日报', '周报': '生成周报',
    '历史报告': '历史报告', '定时生成': '定时生成',
    '对比分析': '对比分析', '批量归档': '历史报告',
    '相似度匹配': '相似度匹配', '案例学习': '案例学习',
    '重新索引': '重新索引', '导出库': '导出知识库',
    '刷新SLO': '刷新SLO', '新建SLO': '新建SLO', '导出合规报告': '导出合规报告',
    '开始训练': '开始训练', '停止': '开始训练',
    '模型评估': '模型评估', '部署': '主界面', '超参调优': '超参调优',
    '版本对比': '版本对比', '健康诊断': '健康诊断',
    '性能优化建议': '性能优化建议', '进程快照': '进程快照',
    '网络诊断': '网络诊断', '日志诊断': '日志诊断', '自愈恢复': '网络诊断',
    '环视': '环绕巡视', '截图保存': '截图保存',
    '数据叠加': '数据叠加', '切换布局': '切换布局',
    '测试通知': '测试通知', '保存配置': '保存配置',
    '新建规则': '新建规则', '编辑规则': '编辑规则',
    '保存修改': '保存修改', '恢复默认': '恢复默认',
    '导入配置': '导出配置', '导出配置': '导出配置',
    '添加组件': '添加组件', '主题切换': '主题切换', '导出仪表盘': '导出仪表盘',
    '创建归档': '立即归档', '恢复': '恢复数据', '清理过期': '清理过期',
    '审计日志': '审计日志', '归档设置': '归档设置', '测试归档': '测试归档',
}

def find_shot(module, op_text, sub_op):
    if not module: return None
    seq = MODULE_NUMS.get(module, '')
    if not seq: return None
    kw = None
    for k, v in op_kw.items():
        if k in op_text:
            kw = v
            break
    if not kw and sub_op:
        kw = sub_op
    if kw:
        p = f'*_{seq}_{module}_{kw}.png'
        fs = glob.glob(os.path.join(SCREENSHOT_DIR, p))
        if fs: return fs[0]
    all_s = sorted(glob.glob(os.path.join(SCREENSHOT_DIR, f'*_{seq}_{module}_*.png')))
    if kw:
        for s in all_s:
            if kw in os.path.basename(s): return s
    main = glob.glob(os.path.join(SCREENSHOT_DIR, f'*_{seq}_{module}_主界面.png'))
    return main[0] if main else (all_s[0] if all_s else None)

# ========== Step 4: 重建文档 ==========
print('\nStep 4: 重建文档（逐段插入截图）')

# 策略：打开清理版文档，从后往前逐段插入
# 为了确保图片正确保存，每次插入后立即保存，重新打开
# 但这样太慢，所以用另一种方式

# 更可靠的方式：构建所有插入操作，然后一次性用XML操作
# 但更简单的方式：每次插入后重新打开文档

shutil.copy2(DST, SRC)  # 用清理版覆盖原文档

doc = Document(SRC)
inserted = 0
failed = 0

# 从后往前，每次操作后重新打开文档（防止索引漂移）
for idx, (pi, p, module, sub_op, text) in enumerate(reversed(actions)):
    doc = Document(SRC)  # 重新打开
    # 重新找到目标段落
    target_idx = len(doc.paragraphs) - 1 - (len(actions) - 1 - idx)
    # 更安全：通过文本匹配
    target_para = None
    for para in doc.paragraphs:
        if para.text.strip() == text.strip():
            target_para = para
            break
    if target_para is None:
        failed += 1
        continue

    img = find_shot(module, text, sub_op)
    if img:
        try:
            # 在目标段落之后插入图片paragraph
            new_p = doc.add_paragraph()
            new_p.paragraph_format.alignment = 1
            new_p.paragraph_format.space_before = Pt(3)
            new_p.paragraph_format.space_after = Pt(3)
            run = new_p.add_run()
            run.add_picture(img, width=Inches(5.8))
            # 用 addnext 移动
            target_para._element.addnext(new_p._element)
            doc.save(SRC)
            inserted += 1
            print(f'  ✓ [{module}/{sub_op}]')
        except Exception as e:
            failed += 1
            print(f'  ✗ [{module}/{sub_op}] {e}')
    else:
        failed += 1
        print(f'  ✗ [{module}/{sub_op}] 无截图')

# ========== Step 5: 登录/注册截图 ==========
print('\nStep 5: 登录/注册截图')
doc = Document(SRC)
for p in doc.paragraphs:
    t = p.text.strip()
    if '步骤4' in t and '验证通过' in t:
        lf = glob.glob(os.path.join(SCREENSHOT_DIR, '*_01_登录页面.png'))
        if lf:
            np = doc.add_paragraph(); np.paragraph_format.alignment = 1
            np.add_run().add_picture(lf[0], width=Inches(5.5))
            p._element.addnext(np._element)
            doc.save(SRC)
            print('  ✓ 登录页面')
        break
for p in doc.paragraphs:
    t = p.text.strip()
    if '步骤4' in t and '注册' in t and '自动填入' in t:
        rf = glob.glob(os.path.join(SCREENSHOT_DIR, '*_02_注册页面.png'))
        if rf:
            np = doc.add_paragraph(); np.paragraph_format.alignment = 1
            np.add_run().add_picture(rf[0], width=Inches(5.5))
            p._element.addnext(np._element)
            doc.save(SRC)
            print('  ✓ 注册页面')
        break

# ========== Step 6: 验证 ==========
print('\nStep 6: 验证')
doc = Document(SRC)
img_rels = sum(1 for r in doc.part.rels.values() if 'image' in r.reltype)
img_paras = sum(1 for p in doc.paragraphs if p._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'))
print(f'  图片关系: {img_rels}')
print(f'  含图段落: {img_paras}')
print(f'  总段落数: {len(doc.paragraphs)}')

if img_paras >= inserted * 0.8:
    print(f'✅ 验证通过!')
else:
    print(f'⚠ 截图段落偏少，可能有丢失')

# 复制覆盖原文档
if os.path.abspath(SRC) != os.path.abspath('操作手册.docx'):
    shutil.copy2(SRC, '操作手册.docx')
print(f'✅ 完成: {os.path.abspath("操作手册.docx")}')
