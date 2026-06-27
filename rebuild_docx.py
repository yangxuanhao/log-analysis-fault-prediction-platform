# -*- coding: utf-8 -*-
"""
重建操作手册：逐段配截图，图文紧挨，字体修复
流程：清理旧截图 → 为每步操作逐个插入截图
"""
import os, sys, glob, re
from docx import Document
from docx.shared import Inches, Pt
from lxml import etree

os.chdir(r'D:\gzwj\信息系统运行日志智能分析与故障预测及根因定位平台')
SCREENSHOT_DIR = 'manual_screenshots'
DOC = '操作手册.docx'

# 已知20个模块名
KNOWN_MODULES = [
    '日志采集汇聚', '日志解析引擎', '异常检测识别', '日志联机搜索',
    '故障预测模型', '根因定位分析', '告警管理中心', '服务拓扑发现',
    '智能诊断修复', '态势感知总览', '智能报告生成', '知识图谱构建',
    'SLA合规监控', '模型训练管理', '系统监控诊断', '多维可视化',
    '告警通知配置', '自定义仪表盘', '配置管理中心', '日志审计归档',
]

MODULE_NUMS = {
    '日志采集汇聚': '01', '日志解析引擎': '02', '异常检测识别': '03',
    '日志联机搜索': '04', '故障预测模型': '05', '根因定位分析': '06',
    '告警管理中心': '07', '服务拓扑发现': '08', '智能诊断修复': '09',
    '态势感知总览': '10', '智能报告生成': '11', '知识图谱构建': '12',
    'SLA合规监控': '13', '模型训练管理': '14', '系统监控诊断': '15',
    '多维可视化': '16', '告警通知配置': '17', '自定义仪表盘': '18',
    '配置管理中心': '19', '日志审计归档': '20',
}

# ============ Step 1: 清理旧截图 ============
print('Step 1: 清理旧截图')
doc = Document(DOC)
# 删除所有包含图片的段落
img_indices = []
for i, p in enumerate(doc.paragraphs):
    draws = p._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
    if draws:
        img_indices.append(i)
print(f'  找到 {len(img_indices)} 个截图段落，删除中...')
for idx in reversed(img_indices):
    doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)
doc.save(DOC)
clean_count = len(Document(DOC).paragraphs)
print(f'  清理后段落数: {clean_count}')

# ============ Step 2: 扫描动作段落 ============
print('\nStep 2: 扫描动作段落')
doc = Document(DOC)
actions = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    if not ('点击「' in text or '拖动「' in text or '勾选「' in text):
        continue

    # 找所属模块
    module = ''
    for j in range(i, -1, -1):
        ht = doc.paragraphs[j].text.strip()
        if ht in KNOWN_MODULES:
            module = ht
            break

    # 找子操作名（最近的Heading 3）
    sub_op = ''
    for j in range(i, -1, -1):
        pj = doc.paragraphs[j]
        if pj.style.name == 'Heading 3':
            ht = pj.text.strip()
            if ht.startswith('- '):
                sub_op = ht[2:]
            elif ht.startswith('-'):
                sub_op = ht[1:]
            else:
                sub_op = ht
            break

    actions.append({'idx': i, 'para': doc.paragraphs[i], 'module': module, 'sub_op': sub_op, 'text': text})
    if module:
        print(f'  [{i:4d}] {module}/{sub_op}')

print(f'  共 {len(actions)} 个动作段落')

# ============ Step 3: 查找匹配截图 ============
print('\nStep 3: 匹配截图')

op_keywords = {
    '一键启动全部': '启动采集', '停止全部': '停止采集',
    '采集统计': '采集统计', '采集配置': '采集配置', '重置计数器': '重置计数器',
    '执行解析': '执行解析', '测试解析': '测试解析', '解析统计': '解析统计',
    '错误模式': '错误模式', '导出结果': '导出结果', '导出报告': '导出报告',
    '立即检测': '立即检测', '检测报告': '检测报告', '高级配置': '高级配置',
    '搜索': '搜索', '实时Tail': '实时Tail', '清空': '搜索',
    '语法帮助': '语法帮助', '执行预测': '执行预测',
    '趋势分析': '趋势分析', '风险评估': '风险评分',
    '根因追溯': '根因追溯', '关联分析': '关联分析',
    '故障注入': '故障注入', '影响域': '影响域', '生成因果图': '生成因果图',
    '模拟告警': '模拟告警', '全部确认': '模拟告警', '清理已处理': '清理已处理',
    '告警统计': '告警统计', '静默模式': '静默模式',
    '自动发现拓扑': '自动发现拓扑', '刷新状态': '刷新状态',
    '拓扑分析': '拓扑分析', '导出拓扑数据': '导出拓扑', '导出拓扑': '导出拓扑',
    '智能诊断': '智能诊断', '一键修复': '一键修复',
    '加入知识库': '加入知识库', '批量诊断': '批量诊断',
    '健康检查': '健康检查', '系统自检': '系统自检',
    '性能快照': '性能快照', '组件状态': '组件状态',
    '流量监控': '流量监控', '生成综合报告': '综合报告',
    '日报': '生成日报', '周报': '生成周报',
    '历史报告': '历史报告', '定时生成': '定时生成',
    '对比分析': '对比分析', '批量归档': '历史报告',
    '相似度匹配': '相似度匹配', '案例学习': '案例学习',
    '重新索引': '重新索引', '导出库': '导出知识库',
    '刷新SLO': '刷新SLO', '新建SLO': '新建SLO', '导出合规报告': '导出合规报告',
    '开始训练': '开始训练', '停止': '开始训练',
    '模型评估': '模型评估', '部署': '主界面', '超参调优': '超参调优',
    '版本对比': '版本对比', '健康诊断': '健康诊断',
    '性能优化建议': '性能优化建议', '进程快照': '进程快照',
    '网络诊断': '网络诊断', '日志诊断': '日志诊断', '自愈恢复': '网络诊断',
    '俯仰角': '主界面', '旋转角': '主界面', '缩放': '主界面',
    '环绕巡视': '环绕巡视', '截图保存': '截图保存',
    '数据叠加': '数据叠加', '切换布局': '切换布局',
    '发送测试通知': '测试通知', '保存配置': '保存配置',
    '新建规则': '新建规则', '编辑规则': '编辑规则',
    '新建': '主界面', '保存': '保存仪表盘', '加载': '主界面',
    '主题切换': '主题切换', '导出仪表盘': '导出仪表盘',
    '添加组件': '添加组件', '保存修改': '保存修改', '恢复默认': '恢复默认',
    '导入配置': '导出配置', '导出配置': '导出配置',
    '创建归档': '立即归档', '恢复': '恢复数据', '清理过期': '清理过期',
    '审计日志': '审计日志', '归档设置': '归档设置', '测试归档': '测试归档',
    '设置保留期限': '主界面',
}

def match_shot(module, op_text, sub_op):
    """为操作文本匹配截图"""
    if not module:
        return None
    seq = MODULE_NUMS.get(module, '')
    if not seq:
        return None

    # 从操作文本提取关键词
    kw = None
    for key, val in op_keywords.items():
        if key in op_text:
            kw = val
            break

    if not kw and sub_op:
        kw = sub_op

    if kw:
        # 精确匹配
        pattern = f'*_{seq}_{module}_{kw}.png'
        files = glob.glob(os.path.join(SCREENSHOT_DIR, pattern))
        if files:
            return files[0]

    # 模糊匹配：在模块截图中找包含关键词的
    all_shots = sorted(glob.glob(os.path.join(SCREENSHOT_DIR, f'*_{seq}_{module}_*.png')))
    if kw:
        for s in all_shots:
            if kw in os.path.basename(s):
                return s

    # 兜底：用主界面
    main = glob.glob(os.path.join(SCREENSHOT_DIR, f'*_{seq}_{module}_主界面.png'))
    if main:
        return main[0]
    # 兜底：用模块第一张
    if all_shots:
        return all_shots[0]
    return None

def insert_pic(doc, after_para, img_path, width=Inches(5.8)):
    """在指定paragraph后面插入截图paragraph"""
    new_p = doc.add_paragraph()
    new_p.paragraph_format.alignment = 1
    new_p.paragraph_format.space_before = Pt(3)
    new_p.paragraph_format.space_after = Pt(3)
    run = new_p.add_run()
    run.add_picture(img_path, width=width)
    after_para._element.addnext(new_p._element)

# ============ Step 4: 插入截图 ============
print('\nStep 4: 逐个插入截图')
doc = Document(DOC)

ok = 0
fail = 0
for act in reversed(actions):
    img = match_shot(act['module'], act['text'], act['sub_op'])
    if img:
        insert_pic(doc, act['para'], img)
        ok += 1
        print(f'  ✓ [{act["module"]}/{act["sub_op"]}] {os.path.basename(img)}')
    else:
        fail += 1
        print(f'  ✗ [{act["module"]}/{act["sub_op"]}] 无截图')

print(f'\n插入完成: 成功 {ok}, 失败 {fail}')

# ============ Step 5: 登录/注册截图 ============
print('\nStep 5: 登录/注册截图')
doc = Document(DOC)
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if '步骤4' in t and '验证通过' in t:
        login = glob.glob(os.path.join(SCREENSHOT_DIR, '*_01_登录页面.png'))
        if login:
            insert_pic(doc, p, login[0])
            print('  ✓ 登录页面')
        break
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if '步骤4' in t and '注册' in t and '自动填入' in t:
        reg = glob.glob(os.path.join(SCREENSHOT_DIR, '*_02_注册页面.png'))
        if reg:
            insert_pic(doc, p, reg[0])
            print('  ✓ 注册页面')
        break

# ============ Step 6: 保存 ============
print('\nStep 6: 保存')
doc.save(DOC)
final = Document(DOC)
print(f'  最终段落数: {len(final.paragraphs)}')
print(f'✅ 已保存: {os.path.abspath(DOC)}')
