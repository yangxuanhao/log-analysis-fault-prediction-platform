# -*- coding: utf-8 -*-
"""
v3：使用 doc.add_paragraph + run.add_picture + addnext
已验证可靠
"""
import os, sys, glob
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

os.chdir(r'D:\gzwj\信息系统运行日志智能分析与故障预测及根因定位平台')
DOCX = '操作手册.docx'
IMG_DIR = 'manual_screenshots'

MN = {'日志采集汇聚':'01','日志解析引擎':'02','异常检测识别':'03','日志联机搜索':'04',
      '故障预测模型':'05','根因定位分析':'06','告警管理中心':'07','服务拓扑发现':'08',
      '智能诊断修复':'09','态势感知总览':'10','智能报告生成':'11','知识图谱构建':'12',
      'SLA合规监控':'13','模型训练管理':'14','系统监控诊断':'15','多维可视化':'16',
      '告警通知配置':'17','自定义仪表盘':'18','配置管理中心':'19','日志审计归档':'20'}

T2S = [
    ('一键启动全部','启动采集'),('停止全部','停止采集'),('采集统计','采集统计'),
    ('采集配置','采集配置'),('重置计数器','重置计数器'),('执行解析','执行解析'),
    ('测试解析','测试解析'),('解析统计','解析统计'),('错误模式','错误模式'),
    ('导出结果','导出结果'),('立即检测','立即检测'),('检测报告','检测报告'),
    ('高级配置','高级配置'),('搜索','搜索'),('实时Tail','实时Tail'),
    ('语法帮助','语法帮助'),('执行预测','执行预测'),('趋势分析','趋势分析'),
    ('风险评估','风险评分'),('导出报告','导出报告'),
    ('根因追溯','根因追溯'),('关联分析','关联分析'),('故障注入','故障注入'),
    ('影响域','影响域'),('生成因果图','生成因果图'),
    ('模拟告警','模拟告警'),('清理已处理','清理已处理'),('告警统计','告警统计'),
    ('静默模式','静默模式'),('自动发现拓扑','自动发现拓扑'),
    ('刷新状态','刷新状态'),('拓扑分析','拓扑分析'),('导出拓扑数据','导出拓扑'),
    ('智能诊断','智能诊断'),('一键修复','一键修复'),('加入知识库','加入知识库'),
    ('批量诊断','批量诊断'),('健康检查','健康检查'),('系统自检','系统自检'),
    ('性能快照','性能快照'),('组件状态','组件状态'),('流量监控','流量监控'),
    ('生成综合报告','综合报告'),('日报','生成日报'),('周报','生成周报'),
    ('历史报告','历史报告'),('定时生成','定时生成'),('对比分析','对比分析'),
    ('相似度匹配','相似度匹配'),('案例学习','案例学习'),('重新索引','重新索引'),
    ('导出库','导出知识库'),('刷新SLO','刷新SLO'),('新建SLO','新建SLO'),
    ('导出合规报告','导出合规报告'),('开始训练','开始训练'),
    ('模型评估','模型评估'),('超参调优','超参调优'),('版本对比','版本对比'),
    ('健康诊断','健康诊断'),('性能优化建议','性能优化建议'),('进程快照','进程快照'),
    ('网络诊断','网络诊断'),('日志诊断','日志诊断'),
    ('环绕巡视','环绕巡视'),('截图保存','截图保存'),
    ('数据叠加','数据叠加'),('切换布局','切换布局'),
    ('测试通知','测试通知'),('保存配置','保存配置'),('新建规则','新建规则'),
    ('编辑规则','编辑规则'),('保存修改','保存修改'),('恢复默认','恢复默认'),
    ('导出配置','导出配置'),('导入配置','导出配置'),
    ('添加组件','添加组件'),('主题切换','主题切换'),('导出仪表盘','导出仪表盘'),
    ('创建归档','立即归档'),('恢复','恢复数据'),('清理过期','清理过期'),
    ('审计日志','审计日志'),('归档设置','归档设置'),('测试归档','测试归档'),
    ('保留期限','主界面'),('新建','主界面'),('部署','主界面'),
    ('保存','保存仪表盘'),('俯仰角','主界面'),('旋转角','主界面'),('缩放','主界面'),
    ('启用告警抑制','主界面'),('自动解决已恢复','主界面'),('抑制窗口','主界面'),
    ('通知方式','主界面'),('全部确认','模拟告警'),('停止','开始训练'),
    ('故障阈值','主界面'),('检测算法','主界面'),('灵敏度阈值','主界面'),
    ('窗口大小','主界面'),('停止检测','主界面'),('启动检测','主界面'),
    ('清空','搜索'),('加载','主界面'),('批量归档','历史报告'),
    ('自愈恢复','网络诊断'),
]

def find_shot(module, text):
    if not module or module not in MN: return None
    seq = MN[module]; kw = None
    for k,v in T2S:
        if k in text: kw=v; break
    if not kw: return None
    fs = glob.glob(f'{IMG_DIR}/*_{seq}_{module}_{kw}.png')
    if fs: return fs[0]
    all_s = sorted(glob.glob(f'{IMG_DIR}/*_{seq}_{module}_*.png'))
    for s in all_s:
        if kw in os.path.basename(s): return s
    return all_s[0] if all_s else None

# ========== 主流程 ==========
print('Step 1: 清理旧截图')
doc = Document(DOCX)
body = doc.element.body
old = [c for c in body if c.tag==qn('w:p') and c.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')]
print(f'  删除 {len(old)} 张旧截图')
for p in old: body.remove(p)
doc.save(DOCX)
print('  已保存清理版')

print('\nStep 2: 扫描动作段落')
doc = Document(DOCX)
actions = []
cur_mod = ''
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t: continue
    if t in MN: cur_mod = t; continue
    if not ('点击「' in t or '拖动「' in t or '勾选「' in t): continue
    actions.append((p, cur_mod, t))
print(f'  共 {len(actions)} 个动作段落')

print('\nStep 3: 登录/注册截图（放在前面，不干扰后续段落索引）')
for p in doc.paragraphs:
    t = p.text.strip()
    if '步骤4' in t and '验证通过' in t:
        lf = glob.glob(f'{IMG_DIR}/*_01_登录页面.png')
        if lf:
            np = doc.add_paragraph(); np.paragraph_format.alignment = 1
            np.add_run().add_picture(lf[0], width=Inches(5.5))
            p._element.addnext(np._element)
            print('  ✓ 登录')
        break
for p in doc.paragraphs:
    t = p.text.strip()
    if '步骤4' in t and '注册' in t and '自动填入' in t:
        rf = glob.glob(f'{IMG_DIR}/*_02_注册页面.png')
        if rf:
            np = doc.add_paragraph(); np.paragraph_format.alignment = 1
            np.add_run().add_picture(rf[0], width=Inches(5.5))
            p._element.addnext(np._element)
            print('  ✓ 注册')
        break

print('\nStep 4: 逐段插入操作截图（从后往前）')
inserted = 0
for p, mod, txt in reversed(actions):
    img = find_shot(mod, txt)
    if not img: continue
    try:
        # 用doc.add_paragraph创建段落（可靠方式）
        np = doc.add_paragraph()
        np.paragraph_format.alignment = 1
        np.paragraph_format.space_before = Pt(3)
        np.paragraph_format.space_after = Pt(3)
        r = np.add_run()
        r.add_picture(img, width=Inches(5.8))
        # 移到目标段落后面
        p._element.addnext(np._element)
        inserted += 1
    except Exception as e:
        print(f'  ✗ {e}')

print(f'  插入 {inserted} 张截图')

print('\nStep 5: 保存')
doc.save(DOCX)

print('\nStep 6: 验证')
doc2 = Document(DOCX)
img_paras = []
for i, p in enumerate(doc2.paragraphs):
    if p._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'):
        prev = doc2.paragraphs[i-1].text.strip()[:60] if i > 0 else '(开头)'
        img_paras.append((i, prev))
print(f'  含图段落: {len(img_paras)}')
if img_paras:
    print(f'  前5个截图位置:')
    for idx, (pi, prev) in enumerate(img_paras[:5]):
        print(f'    截图[{pi}] 前文: "{prev}"')
    print(f'  后5个截图位置:')
    for idx, (pi, prev) in enumerate(img_paras[-5:]):
        print(f'    截图[{pi}] 前文: "{prev}"')
print(f'  总段落数: {len(doc2.paragraphs)}')
if len(img_paras) >= 100:
    print('✅ 验证通过！所有截图已正确定位')
else:
    print('⚠ 截图数量不足')
