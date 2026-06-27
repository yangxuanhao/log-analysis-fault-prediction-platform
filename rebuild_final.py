# -*- coding: utf-8 -*-
"""
最终版：使用python-docx API重建文档
先清理旧图，再用docx API添加图片，最后XML重排body
"""
import os, sys, glob, shutil
from lxml import etree
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

os.chdir(r'D:\gzwj\信息系统运行日志智能分析与故障预测及根因定位平台')
DOCX = '操作手册.docx'
IMG_DIR = 'manual_screenshots'

# 模块名 → 编号
MN = {'日志采集汇聚':'01','日志解析引擎':'02','异常检测识别':'03','日志联机搜索':'04',
      '故障预测模型':'05','根因定位分析':'06','告警管理中心':'07','服务拓扑发现':'08',
      '智能诊断修复':'09','态势感知总览':'10','智能报告生成':'11','知识图谱构建':'12',
      'SLA合规监控':'13','模型训练管理':'14','系统监控诊断':'15','多维可视化':'16',
      '告警通知配置':'17','自定义仪表盘':'18','配置管理中心':'19','日志审计归档':'20'}

# 操作文本 → 截图名映射（完整覆盖所有动作段落）
TEXT2SHOT = [
    ('一键启动全部', '启动采集'), ('停止全部', '停止采集'),
    ('采集统计', '采集统计'), ('采集配置', '采集配置'), ('重置计数器', '重置计数器'),
    ('执行解析', '执行解析'), ('测试解析', '测试解析'), ('解析统计', '解析统计'),
    ('错误模式', '错误模式'), ('导出结果', '导出结果'),
    ('停止检测', '主界面'), ('启动检测', '主界面'),
    ('灵敏度阈值', '主界面'), ('检测算法', '主界面'), ('窗口大小', '主界面'),
    ('立即检测', '立即检测'), ('检测报告', '检测报告'), ('高级配置', '高级配置'),
    ('搜索', '搜索'), ('实时Tail', '实时Tail'), ('清空', '搜索'),
    ('语法帮助', '语法帮助'), ('执行预测', '执行预测'),
    ('故障阈值', '主界面'), ('趋势分析', '趋势分析'),
    ('风险评估', '风险评分'), ('导出报告', '导出报告'),
    ('根因追溯', '根因追溯'), ('关联分析', '关联分析'),
    ('故障注入', '故障注入'), ('影响域', '影响域'), ('生成因果图', '生成因果图'),
    ('模拟告警', '模拟告警'), ('启用告警抑制', '主界面'), ('自动解决已恢复', '主界面'),
    ('抑制窗口', '主界面'), ('通知方式', '主界面'),
    ('全部确认', '模拟告警'), ('清理已处理', '清理已处理'),
    ('告警统计', '告警统计'), ('静默模式', '静默模式'),
    ('自动发现拓扑', '自动发现拓扑'), ('刷新状态', '刷新状态'),
    ('拓扑分析', '拓扑分析'), ('导出拓扑数据', '导出拓扑'),
    ('智能诊断', '智能诊断'), ('一键修复', '一键修复'),
    ('加入知识库', '加入知识库'), ('批量诊断', '批量诊断'),
    ('健康检查', '健康检查'), ('系统自检', '系统自检'),
    ('性能快照', '性能快照'), ('组件状态', '组件状态'),
    ('流量监控', '流量监控'), ('生成综合报告', '综合报告'),
    ('日报', '生成日报'), ('周报', '生成周报'),
    ('历史报告', '历史报告'), ('定时生成', '定时生成'),
    ('对比分析', '对比分析'), ('批量归档', '历史报告'),
    ('相似度匹配', '相似度匹配'), ('案例学习', '案例学习'),
    ('重新索引', '重新索引'), ('导出库', '导出知识库'),
    ('刷新SLO', '刷新SLO'), ('新建SLO', '新建SLO'),
    ('导出合规报告', '导出合规报告'),
    ('开始训练', '开始训练'), ('停止', '开始训练'),
    ('模型评估', '模型评估'), ('部署', '主界面'),
    ('超参调优', '超参调优'), ('版本对比', '版本对比'),
    ('健康诊断', '健康诊断'), ('性能优化建议', '性能优化建议'),
    ('进程快照', '进程快照'), ('网络诊断', '网络诊断'),
    ('日志诊断', '日志诊断'), ('自愈恢复', '网络诊断'),
    ('俯仰角', '主界面'), ('旋转角', '主界面'), ('缩放', '主界面'),
    ('环绕巡视', '环绕巡视'), ('截图保存', '截图保存'),
    ('数据叠加', '数据叠加'), ('切换布局', '切换布局'),
    ('测试通知', '测试通知'), ('保存配置', '保存配置'),
    ('新建规则', '新建规则'), ('编辑规则', '编辑规则'),
    ('新建', '主界面'), ('保存', '保存仪表盘'), ('加载', '主界面'),
    ('主题切换', '主题切换'), ('导出仪表盘', '导出仪表盘'),
    ('添加组件', '添加组件'), ('保存修改', '保存修改'),
    ('恢复默认', '恢复默认'), ('导入配置', '导出配置'), ('导出配置', '导出配置'),
    ('保留期限', '主界面'), ('创建归档', '立即归档'),
    ('恢复', '恢复数据'), ('清理过期', '清理过期'),
    ('审计日志', '审计日志'), ('归档设置', '归档设置'), ('测试归档', '测试归档'),
]

def get_module(text):
    """从段落文本反向查找所属模块名"""
    # 需要外部传入上下文，这里通过文本来判断
    for name in MN:
        if name in text:
            return name
    return ''

def find_shot(module, text):
    """查找匹配的截图"""
    if not module or module not in MN:
        return None
    seq = MN[module]
    # 从操作文本找匹配的截图关键词
    kw = None
    for key, val in TEXT2SHOT:
        if key in text:
            kw = val
            break
    if not kw:
        return None
    # 精确匹配
    fs = glob.glob(f'{IMG_DIR}/*_{seq}_{module}_{kw}.png')
    if fs:
        return fs[0]
    # 模糊匹配
    all_s = sorted(glob.glob(f'{IMG_DIR}/*_{seq}_{module}_*.png'))
    for s in all_s:
        if kw in os.path.basename(s):
            return s
    return all_s[0] if all_s else None

# ========== Phase 1: 清理 ==========
print('Phase 1: 清理旧截图')
doc = Document(DOCX)
body = doc.element.body
old = [c for c in body if c.tag==qn('w:p') and c.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')]
print(f'  删除 {len(old)} 张旧截图')
for p in old: body.remove(p)
doc.save(DOCX)

# ========== Phase 2: 识别动作段落和模块 ==========
print('\nPhase 2: 分析文档')
doc = Document(DOCX)
body = doc.element.body
paras = list(body)  # 所有子元素

# 先用python-docx分析段落
action_items = []  # (段落元素, 模块名, 文本)
cur_mod = ''
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t: continue
    if t in MN:
        cur_mod = t
        continue
    if not ('点击「' in t or '拖动「' in t or '勾选「' in t):
        continue
    action_items.append((p._element, cur_mod, t))

print(f'  动作段落: {len(action_items)}')

# ========== Phase 3: 生成图片段落 ==========
print('\nPhase 3: 生成图片')
img_map = {}  # {目标element: [图片element列表]}
for pe, mod, txt in action_items:
    img_path = find_shot(mod, txt)
    if not img_path:
        # 试最后手段：找模块的主界面
        seq = MN.get(mod, '')
        if seq:
            m = glob.glob(f'{IMG_DIR}/*_{seq}_{mod}_主界面.png')
            if m: img_path = m[0]
    if not img_path:
        print(f'  ✗ [{mod}] {txt[:30]}')
        continue
    try:
        np = doc.add_paragraph()
        np.paragraph_format.alignment = 1
        np.paragraph_format.space_before = Pt(3)
        np.paragraph_format.space_after = Pt(3)
        r = np.add_run()
        r.add_picture(img_path, width=Inches(5.8))
        img_map.setdefault(pe, []).append(np._element)
    except Exception as e:
        print(f'  ✗ {e}')

print(f'  生成 {sum(len(v) for v in img_map.values())} 张截图')

# ========== Phase 4: 重排body ==========
print('\nPhase 4: 重排body')
new_order = []
for child in list(body):
    new_order.append(child)
    if child in img_map:
        for img_elem in img_map[child]:
            new_order.append(img_elem)
# 重建
for child in list(body):
    body.remove(child)
for child in new_order:
    body.append(child)

# ========== Phase 5: 登录/注册 ==========
print('\nPhase 5: 登录/注册')
for p in doc.paragraphs:
    t = p.text.strip()
    if '步骤4' in t and '验证通过' in t:
        lf = glob.glob(f'{IMG_DIR}/*_01_登录页面.png')
        if lf:
            np = doc.add_paragraph(); np.paragraph_format.alignment = 1
            np.add_run().add_picture(lf[0], width=Inches(5.5))
            p._element.addnext(np._element)
            print('  ✓ 登录')
        break
for p in doc.paragraphs:
    t = p.text.strip()
    if '步骤4' in t and '注册' in t and '自动填入' in t:
        rf = glob.glob(f'{IMG_DIR}/*_02_注册页面.png')
        if rf:
            np = doc.add_paragraph(); np.paragraph_format.alignment = 1
            np.add_run().add_picture(rf[0], width=Inches(5.5))
            p._element.addnext(np._element)
            print('  ✓ 注册')
        break

# ========== Phase 6: 保存验证 ==========
print('\nPhase 6: 保存')
doc.save(DOCX)
doc2 = Document(DOCX)
img_rels = sum(1 for r in doc2.part.rels.values() if 'image' in r.reltype)
img_paras = sum(1 for p in doc2.paragraphs if p._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'))
print(f'  图片关系: {img_rels}')
print(f'  含图段落: {img_paras}')
print(f'  总段落数: {len(doc2.paragraphs)}')
if img_paras >= len(action_items) * 0.7:
    print('✅ 验证通过！')
else:
    print('⚠ 截图可能不足')
print(f'✅ 保存至: {os.path.abspath(DOCX)}')
